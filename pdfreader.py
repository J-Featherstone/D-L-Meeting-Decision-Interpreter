import pdfplumber as plmr
import re
import shutil
import docx
import os
import PletstrepReader as ps

#Class that will read PDFs and Gather information about them
class pdfInfo:

    #The filePath is just where the pdf is on the computer.
    def __init__(self, filePath):
        self.filePath = filePath
        #data that needs to be extracted from the PDF
        self.location = ""
        self.appNo = ""
        self.ehDecision = ""
        #creates the decisionPhrasesEH list with
        with open('EHDC Decision Phrases.txt', 'r') as f:
            self.decisionPhrasesEH = [line.strip() for line in f]

    #The PDF file in plaintext, ready for processing information
    

    #This will open the pdf specified when initiating the pdfInfo object.
    def PDFtoText(self):
        
        #opens a pdf, selects the first page, then extracts the text
        pdf = plmr.open(self.filePath)
        page = pdf.pages[0]
        self.text = page.extract_text()

    #testing to make sure getInfo works properly
    def printText(self):
        print(self.text)
    
    #Use regular expressions to take the information out of the string. These may all need to be tweaked depending on how the files are layed out. 
    def extractInfo(self):
        #this line will search the text string for a certain term then return all characters after it and before a new line
        #the .group part at the end of the line isolates the string returned by the expression.
        self.appNo = re.search(r'(?<=Hertford Town Council Our Ref: ).*', self.text).group()
        #print(self.appNo)
        #find the line that contains the full address (identified by everything after "AT:")
        self.locationFull = re.search(r'(?<=AT: ).*?\n', self.text).group()
        self.location = re.search(r'.+?(?= Hertford)', self.locationFull).group()
        #print(self.location)
        #find the decision by looking for key terms.
        self.ehDecision = self.extractDecision()
        #add to the decision list for extraction
        self.decisionList = [self.appNo, self.location, self.ehDecision]

    #the job of extracting the decision from the document is trickier, so I have made an extra function for it.
    #Will search for phrases from a list in a text document that can be updated.
    def extractDecision(self):
        #print(self.decisionPhrasesEH)((.|\n)*)
        #print(self.text)
        txtUpper = self.text.upper()
        txtOneLine = ''.join(txtUpper.splitlines())
        for decisionPhrase in self.decisionPhrasesEH:
            if decisionPhrase.upper() in txtOneLine:
                return decisionPhrase
            #In case the phrase is not in the TXT document, you will need to add it manually
        #self.ehDecision = "Search East Herts Decision Manually"

    #This will run the above functions so you need only call one. Returns a list with the appNo, location and east herts Decision
    def getInfo(self):
        self.PDFtoText()
        #self.printText()
        self.extractInfo()
        return self.decisionList

#A class to allow formatting for multiple pdfs in a certain folder. folderPath is a directory in a string with the pdfs in them
class pdfFolder:
    def __init__(self, meetingDate):
        self.folderPath = r'F:/D & L Committee/PLANNING SUB/PLANS/' + meetingDate + '/Decisions ' + meetingDate + "/"
        self.firstDecisionsList = []
        self.meetingDate = meetingDate
    
    #this function will iterate through the pdf files in the given folder
    def getInfoFromFolder(self):
        for file in os.listdir(self.folderPath):
            print(file)
            if not file.endswith(".pdf"):
                continue
            filePath = self.folderPath + file
            print(filePath)
            pdfFile = pdfInfo(filePath)
            self.firstDecisionsList.append(pdfFile.getInfo())

        pletstrep1 = ps.pletstrep(self.firstDecisionsList, self.meetingDate)
        return pletstrep1.getHTC()


#class for setting up and entering information into a table in a word document.
class createDocx:
    #filepath is the location of the word template.
    #meetingDate will be a string in the format "25.03.2022" - will convert into "25 March 2022"
    def __init__(self, allDecisionList, meetingDate):
        self.allDecisionList = allDecisionList
        self.meetingDate = meetingDate
        self.months = ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"]
        self.formatMeetingDate()
        #self.filePath = filePath + self.fileName

    #Formats the meeting date "25.03.2050" into the format "25 March 2050"
    def formatMeetingDate(self):
        dateList = self.meetingDate.split(".")
        month = self.months[int(dateList[1]) - 1]
        self.titleMeetingDate = dateList[0] + " " + month + " " + dateList[2]
        self.fileName = r'Paper D ' + self.titleMeetingDate + r'.docx'
    
#--------------------------------------------------------------------------------------------------------------------------#
#FilePaths will need to be altered for live use
    def copyTemplate(self):
        original = r'Paper D Template.docx'
        target = r'C:/Users/JoeFeatherstone/Documents/Python Projects/D&L Meeting Decision Interpreter/Test documents/' + self.fileName
        shutil.copyfile(original, target)
        self.filePath = target

    def changeDate(self):
        self.doc = docx.Document(self.filePath)
        for p in self.doc.paragraphs:
            if '*DATE*' in p.text:
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if '*DATE*' in inline[i].text:
                        text = inline[i].text.replace('*DATE*', self.titMeetingDate)
                        inline[i].text = text
                #print(p.text)
                self.doc.save(self.filePath)

    def appendTable(self):
        self.doc.tables
        #print("Retrieved value: " + self.doc.tables[0].cell(0, 0).text)
        #self.doc.tables[0].cell(1, 0).text = "new value1"
        #self.doc.tables[0].add_row() #ADD ROW HERE
        #self.doc.tables[0].cell(2, 1).text = "new value2"
        #self.doc.tables[0].add_row()
        #self.doc.tables[0].cell(3, 2).text = "new value3"
        row = 1
        column = 0
        for decisionList in self.allDecisionList:
            for s in decisionList:
                self.doc.tables[0].cell(row, column).text = s
                column += 1
            row += 1
            column = 0
            self.doc.tables[0].add_row()
        self.doc.save(self.filePath)

#this class is for creating the folders where the PDFS are to be stored
#date in format "dd.mm.yyyy" and initial location is where the folders need to be created
class folders:
    def __init__(self, date, initialLocation):
        self.date = date
        self.initialLocation = initialLocation
        self.locationFolderDate = self.initialLocation + "/" + self.date
    
    def makeFolders(self):
        if not os.path.exists(self.locationFolderDate):
            os.mkdir(self.locationFolderDate)
            os.mkdir(self.locationFolderDate + "/Decisions")
            os.mkdir(self.locationFolderDate + "/Consultations")
            os.mkdir(self.locationFolderDate + "/Paper C")
        else:
            print("Folder already exists")

def getDate():
    Date = input("Enter date of meeting in the format dd.mm.yyyy: ")
    matched = re.match("[0-3][0-9].[0-1][0-9].[0-9][0-9][0-9][0-9]", Date)
    while bool(matched) == False:
        Date = input("Please enter a valid date in the format dd.mm.yyyy: ")
        matched = re.match("[0-3][0-9].[0-1][0-9].[0-9][0-9][0-9][0-9]", Date)
    return Date

#This will wait until the user has moved the relevant PDFS into the folder then entered "Y"
#hopefully this will be replaced with a script that gets the pdfs from the emails automatically
def waitForPDFS():
    confirm = input("Please enter >> y << when pdfs have been moved to the relevant decisions folder: ")
    while confirm != "y":
        confirm = input("Please enter >> y << when pdfs have been moved to the relevant decisions folder: ")



#pdf1 = pdfInfo(r'C:/Users/JoeFeatherstone/Documents/Python Projects/D&L Meeting Decision Interpreter/304A Ware Road.pdf')
#pdf1.printText()
#List1 = pdf1.getInfo()
#print(List1)
#filePath = r'C:/Users/JoeFeatherstone/Documents/Python Projects/D&L Meeting Decision Interpreter/Test documents/'

initialLocation = "F:/D & L Committee/PLANNING SUB/PLANS"
date = getDate()
newFolders = folders(date, initialLocation)
newFolders.makeFolders()
waitForPDFS()
#folder = pdfFolder(date)
#allDecisionsList = folder.getInfoFromFolder()
#print(allDecisionsList)
#paperD = createDocx(allDecisionsList, date)
#paperD.copyTemplate()23
#paperD.changeDate()
#paperD.appendTable()

