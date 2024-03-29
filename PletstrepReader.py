import docx

class pletstrep:
    #appList is the list containing all the current application after extracting them from the PDFs. AppNo is the first entry in each sublist
    #[[appNo, location, EHDecision, HTCDecision], ...] where HTCDesicion will be filled in here.
    #date of meeting is dd.mm.yyyy
    def __init__(self, appList, dateOfMeeting):
        self.dateOfMeeting = dateOfMeeting
        self.appList = appList
        self.HTCDecisions = []

    def getYears(self):
        self.currentYear = self.dateOfMeeting[-4:]
        self.lastYear = str(int(self.currentYear) - 1)

    #this will get the filenapaths of the pletstreps to open (only works on the current filesystem) May need to change hardcoded options
    #This will open the files and save them to self.doc1 and self.doc2
    def openDocs(self):
        self.getYears()
        self.latestFile = "Pletstrep " + self.currentYear + ".docx"
        self.lastYearFile = "Pletstrep " + self.lastYear + ".docx"
        
        #Live documents, be careful!
        #self.latestFilePath = "F:/D & L Committee/PLANNING SUB/PLETRESP/" + self.latestFile
        #self.lastYearFilePath = "F:/D & L Committee/PLANNING SUB/PLETRESP/" + self.lastYearFile

        self.latestFilePath = r"F:/D & L Committee/PLANNING SUB/PLETRESP/" + self.latestFile
        self.lastYearFilePath = r"F:/D & L Committee/PLANNING SUB/PLETRESP/" + self.lastYearFile

        print(self.latestFilePath)
        
        self.doc1 = docx.Document(self.latestFilePath)
        self.doc2 = docx.Document(self.lastYearFilePath)

    #last year is true or false, old version of 
    def iterateTables2(self, document, lastYear):
        print(document.tables[0].cell(0, 0).text)
        for count, application in enumerate(self.appList):
            #debug code
            #print(count)
            appNo = application[0]
            for table1 in document.tables:
                if appNo in table1.cell(0, 0).text:
                    HTCDecision = table1.cell(1, 0).text
                    print(HTCDecision)
                    print(len(self.appList[count]))
                    if len(self.appList[count]) < 4:
                        self.appList[count].append("")
                    #debug code
                    print(self.appList)
                    self.appList[count][3] = self.processDecision(HTCDecision)
                    return True
            if lastYear == True:
                self.appList[count][3] = "No Decision found in last two Pletstrep documents"
        return False

    def iterateTables(self, pletstrepDoc, application):
        appNo = application[0]
        for table1 in pletstrepDoc.tables:
            if appNo in table1.cell(0, 0).text:
                HTCDecision = table1.cell(1, 0).text
                return self.processDecision(HTCDecision)
        return False
    
    
    def processDecision(self, decision):
        decision = decision.lower()
        if not decision:
            return "No Comment"
        elif decision == "no comment":
            return "No Comment"
        elif "no objection" in decision:
            return "No Objection"
        elif "objection" in decision:
            return "Objection"
        else:
            return "Comment, please confirm manually."


    #directory is where the pletstrep word documents are held, it will only go back 1 year to search for applications.
    #This will return the appList with the HTC decision in there as well.
    def getHTC(self):
        self.openDocs()
        
        for count, application in enumerate(self.appList):
            HTCDecision = self.iterateTables(self.doc1, application)
            if HTCDecision == False:
                HTCDecision = self.iterateTables(self.doc2, application)
                if HTCDecision == False:
                    application.append("No Decision found in last two Pletstrep documents")
                else:
                    application.append(HTCDecision)
            else:
                application.append(HTCDecision)
            self.appList[count] = application
                
        return self.appList

#pletstrep1 = pletstrep([["3/21/2739/PNHH", "", "", ""]], "25.04.2022")

#pletstrep1.getHTC()
#print(pletstrep1.appList)